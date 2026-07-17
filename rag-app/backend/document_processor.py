import os
import tempfile
from pathlib import Path
from typing import List

from langchain_core.documents import Document
from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter


SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".docx", ".csv", ".md"}


def load_document(file_bytes: bytes, filename: str) -> List[Document]:
    extension = Path(filename).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise ValueError(f"Unsupported file type '{extension}'. Supported types: {supported}")

    with tempfile.NamedTemporaryFile(delete=False, suffix=extension, mode="wb") as tmp:
        tmp.write(file_bytes)
        tmp_path = tmp.name

    try:
        return _load_by_extension(tmp_path, extension, filename)
    finally:
        os.unlink(tmp_path)


def split_documents(
    documents: List[Document],
    chunk_size: int = 1000,
    chunk_overlap: int = 150,
) -> List[Document]:
    document_offsets = []
    offset = 0
    for document_index, document in enumerate(documents):
        document.metadata["_document_index"] = document_index
        document_offsets.append(offset)
        offset += len(document.page_content) + 2

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", ". ", " ", ""],
    )
    chunks = splitter.split_documents(documents)
    cursors = [0] * len(documents)
    for index, chunk in enumerate(chunks):
        document_index = int(chunk.metadata.pop("_document_index", 0))
        content = documents[document_index].page_content
        search_start = max(0, cursors[document_index] - chunk_overlap - 20)
        local_start = content.find(chunk.page_content, search_start)
        if local_start < 0:
            local_start = content.find(chunk.page_content)
        if local_start < 0:
            local_start = cursors[document_index]
        local_end = local_start + len(chunk.page_content)
        cursors[document_index] = local_end

        chunk.metadata["chunk_index"] = index
        chunk.metadata["char_start"] = document_offsets[document_index] + local_start
        chunk.metadata["char_end"] = document_offsets[document_index] + local_end
    return chunks


def _load_by_extension(path: str, extension: str, original_name: str) -> List[Document]:
    if extension == ".pdf":
        docs = PyPDFLoader(path).load()
        for doc in docs:
            doc.metadata["source"] = original_name
        return docs

    if extension in {".txt", ".md"}:
        content = Path(path).read_text(encoding="utf-8", errors="ignore")
        return [_document(content, original_name)]

    if extension == ".docx":
        from docx import Document as DocxDocument

        doc = DocxDocument(path)
        parts = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return [_document("\n".join(parts), original_name)]

    if extension == ".csv":
        content = Path(path).read_text(encoding="utf-8", errors="ignore")
        return [_document(content, original_name)]

    raise ValueError(f"Unsupported file type: {extension}")


def _document(content: str, source: str) -> Document:
    return Document(page_content=content, metadata={"source": source, "page": 0})
